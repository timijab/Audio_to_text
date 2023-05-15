import docx
from docx.shared import Inches


class WordFormatter:
    def __init__(self, paragraph):
        """ Required parameter: a paragraph or sentence """
        self.initialise = docx.Document()
        self.paragraph = paragraph

    def topic(self):
        """To add headings to the paragraph """
        heading_topic = input("Enter message title: ")
        self.initialise.add_heading(heading_topic, 0)

    def final_paragraph(self):
        """ Adding new paragraph """
        self.initialise.add_paragraph(self.paragraph)
        # self.initialise.paragraph_format.line_spacing = Inches(0.5)
        self.initialise.save('testfile.docx')